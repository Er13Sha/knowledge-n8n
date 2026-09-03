from __future__ import annotations

import csv
import io
import mimetypes
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

import openpyxl
import pymupdf
import pytesseract
from docx import Document
from PIL import Image


class DocumentPreparationError(RuntimeError):
    pass


def normalize_text(value: str) -> str:
    normalized_lines = []
    for line in value.splitlines():
        normalized_line = re.sub(r"[ \t]+", " ", line).strip()
        if normalized_line:
            normalized_lines.append(normalized_line)
    return "\n".join(normalized_lines).strip()


def split_text(value: str, chunk_size: int, overlap: int) -> list[str]:
    if chunk_size <= 0:
        raise ValueError("chunk_size must be positive")
    if overlap < 0 or overlap >= chunk_size:
        raise ValueError("overlap must be between zero and chunk_size")
    text = normalize_text(value)
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text):
            candidate = text[start:end]
            boundary = max(candidate.rfind("\n"), candidate.rfind(" "))
            if boundary > chunk_size // 2:
                end = start + boundary
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks


def prepare_pdf(path: Path, *, chunk_size: int, chunk_overlap: int, ocr_languages: str,
                ocr_dpi: int, ocr_min_text_chars: int, max_pages: int) -> dict[str, list[dict[str, object]]]:
    pages: list[dict[str, object]] = []
    chunks: list[dict[str, object]] = []
    try:
        document = pymupdf.open(path)
    except Exception as exception:
        raise DocumentPreparationError(f"Не удалось открыть PDF: {exception}") from exception
    try:
        if document.page_count == 0:
            raise DocumentPreparationError("PDF не содержит страниц.")
        if document.page_count > max_pages:
            raise DocumentPreparationError(f"PDF содержит {document.page_count} страниц; разрешено не более {max_pages}.")
        for page_index, page in enumerate(document):
            text = normalize_text(page.get_text("text", sort=True))
            ocr_used = False
            if len(text) < ocr_min_text_chars:
                ocr_text = _ocr_page(page, ocr_languages=ocr_languages, ocr_dpi=ocr_dpi)
                if len(ocr_text) > len(text):
                    text = ocr_text
                    ocr_used = True
            page_number = page_index + 1
            pages.append({"page": page_number, "text": text, "ocr_used": ocr_used})
            for page_chunk in split_text(text, chunk_size, chunk_overlap):
                chunks.append({"page": page_number, "chunk_index": len(chunks), "text": page_chunk})
    finally:
        document.close()
    if not chunks:
        raise DocumentPreparationError("В PDF нет текста, который удалось извлечь или распознать.")
    return {"pages": pages, "chunks": chunks}


def extract_document(path: Path, *, filename: str, content_type: str, ocr_languages: str,
                     ocr_dpi: int, max_text_chars: int, max_table_rows: int, max_pages: int) -> dict[str, object]:
    document_format = detect_format(path, filename, content_type)
    if document_format == "pdf":
        return _extract_pdf(path, ocr_languages, ocr_dpi, max_text_chars, max_pages)
    if document_format == "image":
        return _extract_image(path, ocr_languages, max_text_chars)
    if document_format == "docx":
        return _extract_docx(path, max_text_chars)
    if document_format == "doc":
        return _extract_converted_office(path, "doc", max_text_chars)
    if document_format == "xlsx":
        return _extract_xlsx(path, max_text_chars, max_table_rows)
    if document_format == "xls":
        return _extract_converted_office(path, "xls", max_text_chars, max_table_rows)
    if document_format == "csv":
        return _extract_csv(path, max_text_chars, max_table_rows)
    if document_format == "txt":
        return _extract_text(path, max_text_chars)
    raise DocumentPreparationError("Формат файла не поддерживается.")


def detect_format(path: Path, filename: str, content_type: str = "") -> str:
    data = path.read_bytes()[:8192]
    lower_name = filename.lower()
    extension = Path(lower_name).suffix
    if data.startswith(b"%PDF-"):
        return "pdf"
    if data[:4] == b"\xD0\xCF\x11\xE0":
        return "xls" if extension == ".xls" else "doc"
    if data.startswith(b"PK"):
        try:
            with zipfile.ZipFile(path) as archive:
                names = set(archive.namelist())
                content_types = archive.read("[Content_Types].xml").decode("utf-8", "ignore")
                if "word/" in " ".join(names) or "wordprocessingml.document" in content_types:
                    return "docx"
                if "xl/" in " ".join(names) or "spreadsheetml.sheet" in content_types:
                    return "xlsx"
        except (OSError, KeyError, zipfile.BadZipFile):
            pass
    if content_type.startswith("image/") or extension in {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}:
        try:
            with Image.open(path) as image:
                image.verify()
            return "image"
        except Exception as exception:
            raise DocumentPreparationError(f"Изображение повреждено или имеет неизвестный формат: {exception}") from exception
    if extension in {".doc", ".docx", ".xls", ".xlsx", ".csv", ".txt"}:
        return extension[1:]
    decoded = _decode_bytes(data)
    if "," in decoded or ";" in decoded or "\t" in decoded:
        return "csv"
    if decoded.strip():
        return "txt"
    return "other"


def _extract_pdf(path: Path, languages: str, dpi: int, max_chars: int, max_pages: int) -> dict[str, object]:
    try:
        document = pymupdf.open(path)
    except Exception as exception:
        raise DocumentPreparationError(f"Не удалось открыть PDF: {exception}") from exception
    pages: list[str] = []
    ocr_used = False
    metadata = {key.lower(): str(value) for key, value in (document.metadata or {}).items() if value}
    if document.page_count > max_pages:
        document.close()
        raise DocumentPreparationError(f"PDF содержит {document.page_count} страниц; разрешено не более {max_pages}.")
    try:
        for page in document:
            text = normalize_text(page.get_text("text", sort=True))
            if len(text) < 40:
                ocr_text = _ocr_page(page, ocr_languages=languages, ocr_dpi=dpi)
                if len(ocr_text) > len(text):
                    text = ocr_text
                    ocr_used = True
            pages.append(text)
    finally:
        document.close()
    raw_text = "\n\n".join(pages)
    return {"format": "pdf", "mime_type": "application/pdf", "text": _limit(raw_text, max_chars), "pages": len(pages), "tables": [], "metadata": metadata, "stats": {"ocr_used": ocr_used, "pages": len(pages), "text_truncated": len(raw_text) > max_chars}}


def _extract_image(path: Path, languages: str, max_chars: int) -> dict[str, object]:
    try:
        with Image.open(path) as image:
            text = pytesseract.image_to_string(image, lang=languages, config="--psm 3")
    except Exception as exception:
        raise DocumentPreparationError(f"Не удалось распознать изображение: {exception}") from exception
    return {"format": "image", "mime_type": mimetypes.guess_type(path.name)[0] or "image/*", "text": _limit(normalize_text(text), max_chars), "pages": 1, "tables": [], "metadata": {}, "stats": {"ocr_used": True, "pages": 1}}


def _extract_docx(path: Path, max_chars: int) -> dict[str, object]:
    try:
        document = Document(path)
    except Exception as exception:
        raise DocumentPreparationError(f"Не удалось открыть DOCX: {exception}") from exception
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = [_table_from_docx(table, str(index + 1)) for index, table in enumerate(document.tables)]
    table_text = "\n".join("\n".join("\t".join(row) for row in table["rows"]) for table in tables)
    text = normalize_text("\n".join(paragraphs) + "\n" + table_text)
    return {"format": "docx", "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text": _limit(text, max_chars), "pages": 0, "tables": tables, "metadata": {}, "stats": {"tables": len(tables), "rows": sum(len(table["rows"]) for table in tables)}}


def _extract_xlsx(path: Path, max_chars: int, max_table_rows: int) -> dict[str, object]:
    try:
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception as exception:
        raise DocumentPreparationError(f"Не удалось открыть XLSX: {exception}") from exception
    tables = []
    try:
        for sheet in workbook.worksheets:
            rows = [[_cell(value) for value in row] for row in sheet.iter_rows(values_only=True)]
            tables.append(_table_from_rows(rows, sheet.title, max_table_rows))
    finally:
        workbook.close()
    return _spreadsheet_result("xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", tables, max_chars)


def _extract_csv(path: Path, max_chars: int, max_table_rows: int) -> dict[str, object]:
    raw = path.read_bytes()
    text, encoding = _decode_bytes_with_encoding(raw)
    if "\ufffd" in text:
        raise DocumentPreparationError(
            "CSV уже содержит повреждённые символы кодировки (�). "
            "Сохраните исходный файл заново в UTF-8 или Windows-1251 и повторите импорт."
        )
    try:
        dialect = csv.Sniffer().sniff(text[:8192], delimiters=",;\t")
    except csv.Error:
        dialect = csv.excel
    rows = [[str(cell).strip() for cell in row] for row in csv.reader(io.StringIO(text), dialect)]
    table = _table_from_rows(rows, "CSV", max_table_rows)
    result = _spreadsheet_result("csv", "text/csv", [table], max_chars)
    result["encoding"] = encoding
    result["metadata"] = {"encoding": encoding}
    return result


def _extract_text(path: Path, max_chars: int) -> dict[str, object]:
    text = normalize_text(_decode_bytes(path.read_bytes()))
    return {"format": "txt", "mime_type": "text/plain", "text": _limit(text, max_chars), "pages": 0, "tables": [], "metadata": {}, "stats": {}}


def _extract_converted_office(path: Path, original_format: str, max_chars: int, max_table_rows: int = 2000) -> dict[str, object]:
    soffice = shutil.which("soffice")
    if soffice is None:
        raise DocumentPreparationError(f"Для формата {original_format.upper()} не установлен LibreOffice.")
    with tempfile.TemporaryDirectory(prefix="office-") as directory:
        output = Path(directory)
        command = [soffice, "--headless", "--convert-to", "xlsx" if original_format == "xls" else "docx", "--outdir", str(output), str(path)]
        completed = subprocess.run(command, capture_output=True, text=True, timeout=120)
        converted = output / (path.stem + (".xlsx" if original_format == "xls" else ".docx"))
        if completed.returncode != 0 or not converted.exists():
            raise DocumentPreparationError(f"Не удалось преобразовать {original_format.upper()} в читаемый формат.")
        return _extract_xlsx(converted, max_chars, max_table_rows) if original_format == "xls" else _extract_docx(converted, max_chars)


def _spreadsheet_result(document_format: str, mime_type: str, tables: list[dict[str, object]], max_chars: int) -> dict[str, object]:
    blocks = []
    for table in tables:
        blocks.append(f"Лист: {table['sheet']}\n" + "\n".join("\t".join(row) for row in table["rows"]))
    rows = sum(len(table["rows"]) for table in tables)
    return {"format": document_format, "mime_type": mime_type, "text": _limit(normalize_text("\n\n".join(blocks)), max_chars), "pages": 0, "tables": tables, "metadata": {}, "stats": {"sheets": len(tables), "rows": rows}}


def _table_from_docx(table: object, sheet: str) -> dict[str, object]:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    return _table_from_rows(rows, sheet, 2000)


def _table_from_rows(rows: list[list[str]], sheet: str, max_rows: int) -> dict[str, object]:
    rows = [row for row in rows if any(cell.strip() for cell in row)]
    width = max((len(row) for row in rows), default=0)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    columns = normalized[0] if normalized else []
    data = normalized[1:] if normalized else []
    truncated = len(data) > max_rows
    return {"sheet": sheet, "columns": columns, "rows": data[:max_rows], "truncated": truncated}


def _cell(value: object) -> str:
    return "" if value is None else str(value)


def _decode_bytes(value: bytes) -> str:
    return _decode_bytes_with_encoding(value)[0]


def _decode_bytes_with_encoding(value: bytes) -> tuple[str, str]:
    """Decode common document encodings without turning UTF-16 into mojibake."""
    encodings = (
        (b"\xff\xfe\x00\x00", "utf-32-le"),
        (b"\x00\x00\xfe\xff", "utf-32-be"),
        (b"\xff\xfe", "utf-16-le"),
        (b"\xfe\xff", "utf-16-be"),
        (b"\xef\xbb\xbf", "utf-8-sig"),
    )
    for bom, encoding in encodings:
        if value.startswith(bom):
            return value.decode(encoding).lstrip("\ufeff"), encoding

    if _looks_like_utf16(value):
        for encoding in ("utf-16-le", "utf-16-be"):
            try:
                decoded = value.decode(encoding)
            except UnicodeDecodeError:
                continue
            if _is_readable(decoded):
                return decoded, encoding

    for encoding in ("utf-8", "cp1251", "cp866", "koi8-r", "windows-1252", "latin-1"):
        try:
            decoded = value.decode(encoding)
        except UnicodeDecodeError:
            continue
        if _is_readable(decoded):
            return decoded, encoding
    return value.decode("utf-8", errors="replace"), "utf-8-replacement"


def _looks_like_utf16(value: bytes) -> bool:
    sample = value[:4096]
    if len(sample) < 8:
        return False
    zero_bytes = sample.count(0)
    return zero_bytes / len(sample) > 0.2


def _is_readable(value: str) -> bool:
    if not value:
        return True
    control_characters = sum(1 for character in value if ord(character) < 32 and character not in "\r\n\t")
    return control_characters / len(value) < 0.02


def _limit(value: str, limit: int) -> str:
    return value[:max(1, limit)]


def _ocr_page(page: pymupdf.Page, *, ocr_languages: str, ocr_dpi: int) -> str:
    try:
        pixmap = page.get_pixmap(dpi=ocr_dpi, alpha=False)
        image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
        return normalize_text(pytesseract.image_to_string(image, lang=ocr_languages))
    except Exception as exception:
        raise DocumentPreparationError(f"Не удалось распознать страницу {page.number + 1}: {exception}") from exception
